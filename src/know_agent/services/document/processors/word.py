"""Word 处理器 — python-docx.

替代源项目的 MinerU 调用。仅支持 .docx（旧版 .doc 需先转换）。
"""

from io import BytesIO

from docx import Document

from know_agent.models.enums import FileType, KnowledgeBaseType
from know_agent.services.document.processors.base import FileProcessor


class WordProcessor(FileProcessor):
    def supports(self, file_type: FileType, kb_type: KnowledgeBaseType | None) -> bool:
        return file_type == FileType.DOC

    def parse(self, content: bytes) -> str:
        doc = Document(BytesIO(content))
        parts: list[str] = [p.text for p in doc.paragraphs if p.text]
        for table in doc.tables:
            for row in table.rows:
                cells = [c.text.strip() for c in row.cells]
                parts.append(" | ".join(cells))
        return "\n\n".join(parts)
